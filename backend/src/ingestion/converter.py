import os
import json
import pandas as pd
from bs4 import BeautifulSoup
from datetime import datetime
import logging
from functools import wraps

# Set up logging for conversion
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
if not logger.handlers:
    ch = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    ch.setFormatter(formatter)
    logger.addHandler(ch)

class DataConverter:
    """
    Member 1's Task: The Converter.
    Unifies disparate raw formats (JSON, HTML, etc) into a single Structured CSV and Parquet Hub.
    """
    def __init__(self):
        # Discover base directory
        self.base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.raw_dir = os.path.join(self.base_dir, "data", "raw")
        self.output_dir = os.path.join(self.base_dir, "data", "processed")
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _get_latest_file(self, prefix, extension):
        """Finds the most recent file for a given source using modification time."""
        if not os.path.exists(self.raw_dir):
            return None
            
        files = [f for f in os.listdir(self.raw_dir) if f.startswith(prefix) and f.endswith(extension)]
        if not files:
            return None
        # Use full path for mtime comparison
        full_paths = [os.path.join(self.raw_dir, f) for f in files]
        return max(full_paths, key=os.path.getmtime)

    def _flatten_dict(self, d, parent_key='', sep='_'):
        """Helper to flatten nested dictionaries for CSV compatibility using iteration to avoid recursion limits."""
        items = []
        stack = [((), d)]
        while stack:
            path, current = stack.pop()
            for k, v in current.items():
                new_key = path + (k,)
                if isinstance(v, dict) and v:
                    stack.append((new_key, v))
                else:
                    items.append((sep.join(new_key), v))
        return dict(items)
        
    def _fetch_file(prefix, extension):
        """Decorator to fetch the latest file automatically and inject it into the parsing method. Returns [] if not found."""
        def decorator(func):
            @wraps(func)
            def wrapper(self, *args, **kwargs):
                file_path = self._get_latest_file(prefix, extension)
                if not file_path:
                    return []
                return func(self, file_path, *args, **kwargs)
            return wrapper
        return decorator

    @_fetch_file("api_data", "json")
    def parse_api_json(self, file_path):
        """Extracts and flattens keys from the API JSON dynamically."""
        logger.info(f"Parsing Dynamic JSON (Flattened): {os.path.basename(file_path)}...")
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
        except Exception as e:
            logger.error(f"API JSON Parsing Error: {e}")
            return []
        
        # Assume 'data' or 'properties' key might contain the list
        records = data.get("data") or data.get("properties") or data
        if not isinstance(records, list): 
            records = [records]

        standardized = []
        for rec in records:
            if isinstance(rec, dict):
                entry = self._flatten_dict(rec)
            else:
                entry = {"value": str(rec)}
            
            entry["source"] = "API_Source"
            entry["ingested_at"] = datetime.now().isoformat()
            standardized.append(entry)
        return standardized

    @_fetch_file("json_upload", "json")
    def parse_json_upload(self, file_path):
        """Extracts and flattens keys from uploaded JSON files."""
        logger.info(f"Parsing JSON Upload (Flattened): {os.path.basename(file_path)}...")
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
        except Exception as e:
            logger.error(f"JSON Upload Parsing Error: {e}")
            return []

        # Try to find the list of records
        records = None
        if isinstance(data, list):
            records = data
        elif isinstance(data, dict):
            # Common keys for data lists
            for key in ['data', 'properties', 'results', 'items', 'records', 'features']:
                if key in data and isinstance(data[key], list):
                    records = data[key]
                    break
            # Fallback: treat the whole dict as one record if no list found
            if records is None:
                records = [data]
        
        if not records:
            records = []
            
        logger.info(f"JSON Upload: Found {len(records)} records.")

        standardized = []
        for rec in records:
            if isinstance(rec, dict):
                entry = self._flatten_dict(rec)
            else:
                entry = {"value": str(rec)}
            
            entry["source"] = "JSON_Upload"
            entry["ingested_at"] = datetime.now().isoformat()
            standardized.append(entry)
        return standardized

    @_fetch_file("xlsx_upload", "xlsx")
    def parse_xlsx_upload(self, file_path):
        """Extracts data from uploaded Excel files using Pandas."""
        logger.info(f"Parsing Excel Upload: {os.path.basename(file_path)}...")
        try:
            # Read the first sheet by default
            df = pd.read_excel(file_path)
            
            # Simple sanitization: replace NaN with empty string
            df = df.fillna("")
            
            standardized = []
            for _, row in df.iterrows():
                entry = row.to_dict()
                entry["source"] = "XLSX_Upload"
                entry["ingested_at"] = datetime.now().isoformat()
                standardized.append(entry)
                
            return standardized
        except Exception as e:
            logger.error(f"XLSX Parsing Error: {e}")
            return []

    @_fetch_file("xml_upload", "xml")
    def parse_xml_upload(self, file_path):
        """Extracts data from uploaded XML files using Pandas read_xml or simple ElementTree."""
        logger.info(f"Parsing XML Upload: {os.path.basename(file_path)}...")
        try:
            # Pandas read_xml is very robust. We explicitly set encoding for robustness.
            df = pd.read_xml(file_path, encoding="utf-8")
            df = df.fillna("")
            
            standardized = []
            for _, row in df.iterrows():
                entry = row.to_dict()
                entry["source"] = "XML_Upload"
                entry["ingested_at"] = datetime.now().isoformat()
                standardized.append(entry)
                
            return standardized
        except Exception as e:
            logger.error(f"XML Parsing Error with pandas (trying fallback): {e}")
            try:
                import xml.etree.ElementTree as ET
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    tree = ET.parse(f)
                root = tree.getroot()
                records = []
                for child in root:
                    record = {}
                    for subchild in child:
                        record[subchild.tag] = str(subchild.text) if subchild.text else ""
                    record["source"] = "XML_Upload"
                    record["ingested_at"] = datetime.now().isoformat()
                    records.append(record)
                return records
            except Exception as e2:
                logger.error(f"XML Fallback Parsing Error: {e2}")
                return []

    @_fetch_file("parquet_upload", "parquet")
    def parse_parquet_upload(self, file_path):
        """Standardized Parquet Parser: Reads uploaded binary columnar data."""
        logger.info(f"Parsing Parquet Upload: {os.path.basename(file_path)}...")
        try:
            df = pd.read_parquet(file_path)
            df = df.fillna("")
            
            standardized = []
            for _, row in df.iterrows():
                entry = row.to_dict()
                entry["source"] = "Parquet_Upload"
                entry["ingested_at"] = datetime.now().isoformat()
                standardized.append(entry)
                
            return standardized
        except Exception as e:
            logger.error(f"Parquet Parsing Error: {e}")
            return []

    @_fetch_file("universal_", "")
    def parse_other_upload(self, file_path):
        """Universal Format Parser (Heuristic Sensing Engine).
        Detects binary formats first, then falls back to text heuristics."""
        logger.info(f"Heuristic Sensing on Universal Format: {os.path.basename(file_path)}...")
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            # Step 0: Binary Format Detection (Excel, Parquet, JSON)
            if ext in ['.xlsx', '.xls']:
                logger.info("Detected Excel binary format. Using pd.read_excel.")
                df = pd.read_excel(file_path)
                df = df.fillna("")
                standardized = []
                for _, row in df.iterrows():
                    entry = row.to_dict()
                    entry["source"] = "Universal_Excel"
                    entry["ingested_at"] = datetime.now().isoformat()
                    standardized.append(entry)
                return standardized
                
            if ext == '.parquet':
                logger.info("Detected Parquet binary format. Using pd.read_parquet.")
                df = pd.read_parquet(file_path)
                df = df.fillna("")
                standardized = []
                for _, row in df.iterrows():
                    entry = row.to_dict()
                    entry["source"] = "Universal_Parquet"
                    entry["ingested_at"] = datetime.now().isoformat()
                    standardized.append(entry)
                return standardized
                
            if ext == '.json':
                logger.info("Detected JSON format. Using pd.read_json.")
                import json as json_lib
                with open(file_path, "r", encoding="utf-8") as f:
                    raw = json_lib.load(f)
                records = raw if isinstance(raw, list) else [raw]
                standardized = []
                for rec in records:
                    entry = self._flatten_dict(rec) if isinstance(rec, dict) else {"value": str(rec)}
                    entry["source"] = "Universal_JSON"
                    entry["ingested_at"] = datetime.now().isoformat()
                    standardized.append(entry)
                return standardized

            # Step 0b: Document Format Detection (PDF, DOCX)
            if ext == '.pdf':
                logger.info("Detected PDF format. Routing to dedicated PDF parser.")
                return self.parse_pdf_document.__wrapped__(self, file_path)
                
            if ext in ['.docx', '.doc']:
                logger.info("Detected DOCX format. Routing to dedicated DOCX parser.")
                return self.parse_docx_document.__wrapped__(self, file_path)

            # Step 1: Text-based Delimiter Sensing
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                head = f.read(2048) # Sample first 2kb
            
            # Simple heuristic: counts common delims
            delims = [',', ';', '|', '\t']
            found_delim = None
            for d in delims:
                if head.count(d) > 10: 
                    found_delim = d
                    break
            
            if found_delim:
                logger.info(f"Detected Table structure with Delim: '{found_delim}'")
                df = pd.read_csv(file_path, sep=found_delim, encoding="utf-8", encoding_errors="replace", on_bad_lines="skip")
                df = df.fillna("")
                standardized = []
                for _, row in df.iterrows():
                    entry = row.to_dict()
                    entry["source"] = "Universal_Heuristic"
                    entry["ingested_at"] = datetime.now().isoformat()
                    standardized.append(entry)
                return standardized

            # Step 2: Key-Value Sensing
            import re
            lines = head.splitlines()
            kv_map = {}
            for line in lines:
                match = re.match(r"^([\w\s_-]+)[:=](.*)$", line.strip())
                if match:
                    key, val = match.groups()
                    kv_map[key.strip().lower().replace(" ", "_")] = val.strip()
            
            if len(kv_map) >= 3:
                logger.info(f"Detected Key-Value structure ({len(kv_map)} fields)")
                entry = kv_map.copy()
                entry["source"] = "Universal_Heuristic_KV"
                entry["sensed_format"] = "Key-Value Pairs"
                entry["filename"] = os.path.basename(file_path)
                entry["ingested_at"] = datetime.now().isoformat()
                return [entry]

            # Step 3: Paragraph/Line Sensing (Default)
            logger.info("Falling back to Paragraph/Line Sensing")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                full_content = f.read().strip()
            
            if not full_content:
                return []

            if "\n\n" in full_content:
                blocks = [b.strip() for b in full_content.split("\n\n") if b.strip()]
            else:
                blocks = [l.strip() for l in full_content.splitlines() if l.strip()]
                
            standardized = []
            for block in blocks:
                standardized.append({
                    "raw_content": block,
                    "filename": os.path.basename(file_path),
                    "source": "Universal_Heuristic_Text",
                    "sensed_format": "Unstructured Text Blocks",
                    "ingested_at": datetime.now().isoformat()
                })
            return standardized

        except Exception as e:
            logger.error(f"Universal Parsing Error: {e}")
            return []

    @_fetch_file("zip_upload", "zip")
    def parse_zip_upload(self, file_path):
        """Extracts and scans multimedia health from uploaded ZIP archives."""
        import zipfile
        from PIL import Image
        import io
        
        logger.info(f"Scanning Multimedia ZIP: {os.path.basename(file_path)}...")
        standardized = []
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                for file_info in zf.infolist():
                    if file_info.is_dir() or "__MACOSX" in file_info.filename or file_info.filename.startswith("._"):
                        continue
                        
                    filename = file_info.filename
                    ext = filename.split('.')[-1].lower() if '.' in filename else ""
                    size_kb = round(file_info.file_size / 1024, 2)
                    
                    record = {
                        "filename": filename,
                        "extension": ext,
                        "size_kb": size_kb,
                        "status": "Healthy",
                        "multimedia_type": "Other",
                        "resolution": "—",
                        "source": "ZIP_Archive",
                        "ingested_at": datetime.now().isoformat()
                    }
                    
                    if ext in ['jpg', 'jpeg', 'png', 'svg', 'webp', 'gif']:
                        record["multimedia_type"] = "Image"
                        try:
                            with zf.open(filename) as f:
                                img_data = f.read()
                                img = Image.open(io.BytesIO(img_data))
                                width, height = img.width, img.height
                                record["resolution"] = f"{width}x{height}"
                                
                                megapixels = (width * height) / 1_000_000
                                if megapixels >= 3.0:
                                    record["suitability"] = "Professional Print"
                                elif megapixels >= 0.5:
                                    record["suitability"] = "AI Training / HD"
                                elif megapixels >= 0.1:
                                    record["suitability"] = "Standard Web"
                                else:
                                    record["suitability"] = "Low Quality"
                                    record["status"] = "Warning (Low Res)"
                        except Exception:
                            record["status"] = "Corrupted/Invalid"
                            record["suitability"] = "Unusable"
                    else:
                        record["suitability"] = "N/A (Non-Image)"
                        
                    standardized.append(record)
                    
            logger.info(f"ZIP Scan Complete: {len(standardized)} items found.")
            return standardized
        except Exception as e:
            logger.error(f"ZIP Parsing Error: {e}")
            return []

    @_fetch_file("web_scrape", "html")
    def parse_city_html(self, file_path):
        """Standardized HTML Table Parser: Extracts data from all tables on the page."""
        logger.info(f"Parsing Dynamic HTML (Standardized): {os.path.basename(file_path)}...")
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            
            all_tables = soup.find_all("table")
            if not all_tables:
                return self._parse_city_html_fallback(soup)

            standardized = []
            for table in all_tables:
                rows = table.find_all("tr")
                if not rows: continue

                header_row = table.find("thead").find("tr") if table.find("thead") else rows[0]
                headers = [th.get_text(strip=True) for th in header_row.find_all(["th", "td"])]
                
                data_rows = rows[1:] if not table.find("thead") else rows
                
                for row in data_rows:
                    cols = row.find_all(["td", "th"])
                    if not cols or len(cols) == 0: continue
                    
                    entry = {}
                    for i in range(min(len(headers), len(cols))):
                        key = headers[i] if headers[i] else f"Col_{i}"
                        entry[key] = cols[i].get_text(strip=True)
                    
                    if entry and any(entry.values()):
                        entry["source"] = "Web_Scrape"
                        entry["ingested_at"] = datetime.now().isoformat()
                        standardized.append(entry)
            
            if not standardized:
                return self._parse_city_html_fallback(soup)
                
            return standardized[:100]
        except Exception as e:
            logger.error(f"HTML Parsing Error: {e}")
            return []

    def _parse_city_html_fallback(self, soup):
        logger.info("No tables found, checking for list items (li)...")
        standardized = []
        lists = soup.find_all(["ul", "ol"])
        for lst in lists:
            items = lst.find_all("li")
            if len(items) < 5: continue
            for li in items:
                text = li.get_text(strip=True)
                if text:
                    standardized.append({
                        "Item_Content": text,
                        "source": "Web_Scrape",
                        "ingested_at": datetime.now().isoformat()
                    })
        return standardized[:100]

    @_fetch_file("user_upload", "csv")
    def parse_user_csv(self, file_path):
        """Extracts all columns from manually uploaded CSVs."""
        logger.info(f"Parsing User CSV: {os.path.basename(file_path)}...")
        try:
            df = pd.read_csv(file_path, encoding="utf-8", encoding_errors="replace")
            df = df.fillna("")
            
            standardized = []
            for _, row in df.iterrows():
                entry = row.to_dict()
                entry["source"] = "User_Upload"
                entry["ingested_at"] = datetime.now().isoformat()
                standardized.append(entry)
                
            return standardized
        except Exception as e:
            logger.error(f"CSV Parsing Error: {e}")
            return []

    @_fetch_file("pdf_upload", "pdf")
    def parse_pdf_document(self, file_path):
        """Extracts structured data (tables, KV pairs) from uploaded PDF documents.
        Uses strict per-page priority: Tables > KV Pairs > Text Lines."""
        logger.info(f"Parsing PDF Document (Structured): {os.path.basename(file_path)}...")
        standardized = []
        try:
            import pdfplumber
            import re
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_records = []
                    
                    # Priority 1: Table Extraction
                    tables = page.extract_tables({
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                        "snap_tolerance": 3,
                    })
                    if not tables:
                        tables = page.extract_tables()
                    
                    for table in tables:
                        if len(table) > 1:
                            headers = [str(h).strip() if h else f"Col_{j}" for j, h in enumerate(table[0])]
                            for row in table[1:]:
                                entry = {headers[j]: (str(val).strip() if val else "") for j, val in enumerate(row) if j < len(headers)}
                                if any(entry.values()):
                                    entry["source"] = "PDF_Table"
                                    entry["page"] = i + 1
                                    entry["ingested_at"] = datetime.now().isoformat()
                                    page_records.append(entry)
                    
                    # If tables were found on this page, skip text parsing
                    if page_records:
                        standardized.extend(page_records)
                        continue
                    
                    # Priority 2: Key-Value Pair Extraction (only if no tables)
                    text = page.extract_text()
                    if text:
                        kv_pairs = re.findall(r'^\s*([\w\s]+)\s*:\s*(.+)$', text, re.MULTILINE)
                        if len(kv_pairs) >= 2:
                            kv_entry = {k.strip().replace(' ', '_'): v.strip() for k, v in kv_pairs}
                            kv_entry["source"] = "PDF_KV_Pairs"
                            kv_entry["page"] = i + 1
                            kv_entry["ingested_at"] = datetime.now().isoformat()
                            standardized.append(kv_entry)
                            continue
                    
                        # Priority 3: Line-by-line fallback (only if nothing else worked)
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        for line_idx, line_content in enumerate(lines):
                            standardized.append({
                                "Content": line_content,
                                "Page_Number": i + 1,
                                "Line_Number": line_idx + 1,
                                "source": "PDF_Text_Line",
                                "ingested_at": datetime.now().isoformat()
                            })
        except ImportError:
            logger.error("pdfplumber is not installed. To parse PDF files, run 'pip install pdfplumber'.")
        except Exception as e:
            logger.error(f"PDF Parsing Error: {e}")
            
        return standardized

    @_fetch_file("docx_upload", "docx")
    def parse_docx_document(self, file_path):
        """Extracts structured data from Word documents.
        Strict priority: Tables > Grouped KV Pairs > Full Text."""
        logger.info(f"Parsing DOCX Document (Structured): {os.path.basename(file_path)}...")
        standardized = []
        try:
            from docx import Document
            import re
            doc = Document(file_path)
            
            # Priority 1: Table Extraction
            for table in doc.tables:
                if len(table.rows) > 1:
                    headers = [cell.text.strip() if cell.text.strip() else f"Col_{j}" for j, cell in enumerate(table.rows[0].cells)]
                    for row in table.rows[1:]:
                        entry = {headers[j]: row.cells[j].text.strip() for j in range(min(len(headers), len(row.cells)))}
                        if any(entry.values()):
                            entry["source"] = "DOCX_Table"
                            entry["ingested_at"] = datetime.now().isoformat()
                            standardized.append(entry)
            
            # If tables were found, skip text parsing entirely
            if standardized:
                return standardized
            
            # Priority 2: Grouped Key-Value Pair Extraction
            kv_map = {}
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    match = re.match(r'^([^:]+)\s*:\s*(.+)$', text)
                    if match:
                        label, val = match.groups()
                        kv_map[label.strip().replace(' ', '_')] = val.strip()
            
            if len(kv_map) >= 2:
                kv_map["source"] = "DOCX_KV_Pairs"
                kv_map["ingested_at"] = datetime.now().isoformat()
                standardized.append(kv_map)
                return standardized
            
            # Priority 3: Full text fallback
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text:
                standardized.append({
                    "Content": full_text,
                    "source": "DOCX_Text",
                    "ingested_at": datetime.now().isoformat()
                })
        except Exception as e:
            logger.error(f"DOCX Parsing Error: {e}")
            
        return standardized

    def unify_to_parquet(self, source_filter=None):
        """Combines all sources and saves to /data/processed/raw_structured.parquet."""
        # Dynamic mapping to avoid explicit IF cascades
        source_map = {
            "api": self.parse_api_json,
            "scraping": self.parse_city_html,
            "upload": self.parse_user_csv,
            "pdf": self.parse_pdf_document,
            "docx": self.parse_docx_document,
            "json_upload": self.parse_json_upload,
            "xlsx_upload": self.parse_xlsx_upload,
            "zip_upload": self.parse_zip_upload,
            "xml_upload": self.parse_xml_upload,
            "parquet_upload": self.parse_parquet_upload,
            "others_upload": self.parse_other_upload
        }

        all_data = []
        
        # Execute only the requested filter parser, or execute all if None
        targets = [source_filter] if source_filter in source_map else source_map.keys()
        for k in targets:
            parser_func = source_map[k]
            extracted = parser_func()
            if extracted:
                all_data.extend(extracted)
        
        if not all_data:
            logger.warning("No data found to convert.")
            return None
        
        df = pd.DataFrame(all_data)
        
        # Security/Format Fix: Force all heterogeneous Excel/JSON/API native objects
        df = df.astype(str)
        
        output_path = os.path.join(self.output_dir, "raw_structured.parquet")
        df.to_parquet(output_path, index=False)
        
        csv_path = os.path.join(self.output_dir, "raw_structured.csv")
        df.to_csv(csv_path, index=False)
        
        logger.info(f"Success! Unified Parquet Hub created at: {output_path}")
        logger.info(f"Secondary CSV Export available at: {csv_path}")
        return output_path

if __name__ == "__main__":
    converter = DataConverter()
    converter.unify_to_parquet()
